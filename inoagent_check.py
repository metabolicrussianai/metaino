#!/usr/bin/env python3
'''inoagent_check — оффлайн-вычитка документов на упоминания лиц и организаций
из реестра иностранных агентов Минюста России.

Не является юридической консультацией. Утилита выдаёт отчёт человеку;
решение о правке текста принимает человек.

Сеть используется только командой update. Команда scan работает оффлайн.
'''
from __future__ import annotations

import argparse
import hashlib
import json
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

STORE = Path.home() / '.inoagent_check'
SNAPSHOT_DIR = STORE / 'snapshots'
MANIFEST = STORE / 'manifest.json'

MIN_RECORDS = 300
MAX_DELTA = 0.25
WINDOW_CHARS = 600
HEAD_CHARS = 800
DEFAULT_MIN_CONFIDENCE = 0.6

STATUS_MARKERS = (
    'иностранного агента',
    'иностранным агентом',
    'иностранный агент',
    'иноагент',
    'реестр иностранных агентов',
)

NBSP = '\u00a0'
SOFT_HYPHEN = '\u00ad'

INDECLINABLE_TAILS = ('ко', 'о', 'е', 'и', 'у', 'ю', 'ых', 'их', 'аго')


def norm(text):
    out = str(text).replace(NBSP, ' ').replace(SOFT_HYPHEN, '')
    return out.lower().replace('ё', 'е')


def tokenize(text):
    '''Список (форма, начало, конец). Без регулярных выражений.'''
    tokens = []
    i = 0
    n = len(text)
    while i < n:
        if text[i].isalpha():
            j = i
            while j < n and (text[j].isalpha() or text[j] == '-'):
                j += 1
            while j > i and not text[j - 1].isalpha():
                j -= 1
            tokens.append((norm(text[i:j]), i, j))
            i = j
        else:
            i += 1
    return tokens


def surname_forms(surname):
    '''Падежные формы фамилии по правилам, без морфологического словаря.

    Редкие и иноязычные фамилии могут не попасть в выборку. Это
    ложноотрицательный результат — задокументированное ограничение,
    а не недосмотр. Критичные персоны проверяются вручную.
    '''
    s = norm(surname)
    forms = {s}
    if len(s) < 3 or s.endswith(INDECLINABLE_TAILS):
        return forms
    if s.endswith(('ова', 'ева', 'ина', 'ына')):
        forms.update(s[:-1] + e for e in ('ой', 'у', 'е'))
    elif s.endswith(('ов', 'ев', 'ин', 'ын')):
        forms.update(s + e for e in ('а', 'у', 'ым', 'ом', 'е'))
    elif s.endswith('ий'):
        forms.update(s[:-2] + e for e in ('ого', 'ому', 'им', 'ом'))
    elif s.endswith('ая'):
        forms.update(s[:-2] + e for e in ('ой', 'ую'))
    elif s.endswith('ой'):
        forms.update(s[:-2] + e for e in ('ого', 'ому', 'ым', 'ом'))
    elif s.endswith('ь'):
        forms.update(s[:-1] + e for e in ('я', 'ю', 'ем', 'е'))
    elif s.endswith('а'):
        forms.update(s[:-1] + e for e in ('ы', 'е', 'у', 'ой'))
    elif s.endswith('я'):
        forms.update(s[:-1] + e for e in ('и', 'ю', 'ей'))
    else:
        forms.update(s + e for e in ('а', 'у', 'ом', 'е'))
    return forms


@dataclass(frozen=True)
class Record:
    rid: int
    name: str
    kind: str
    included_on: str


@dataclass(frozen=True)
class Variant:
    slots: tuple
    confidence: float
    rid: int
    label: str


@dataclass
class Hit:
    rid: int
    label: str
    confidence: float
    start: int
    end: int
    quote: str
    marked: bool = False


def person_variants(rec):
    parts = [p for p, _, _ in tokenize(rec.name)]
    if not parts:
        return []
    head = frozenset(surname_forms(parts[0]))
    rest = [frozenset({p, p[0]}) for p in parts[1:]]
    out = []
    if rest:
        out.append(Variant((head,) + tuple(rest), 0.95, rec.rid, rec.name))
        out.append(Variant((head, rest[0]), 0.88, rec.rid, rec.name))
    alone = 0.72 if len(parts[0]) >= 7 else 0.55
    out.append(Variant((head,), alone, rec.rid, rec.name))
    return out


def org_variants(rec):
    parts = [p for p, _, _ in tokenize(rec.name)]
    if not parts:
        return []
    slots = tuple(frozenset({p}) for p in parts)
    confidence = 0.9 if len(parts) > 1 else 0.75
    return [Variant(slots, confidence, rec.rid, rec.name)]


def build_index(records):
    index = {}
    for rec in records:
        maker = person_variants if rec.kind == 'person' else org_variants
        for var in maker(rec):
            for form in var.slots[0]:
                index.setdefault(form, []).append(var)
    for bucket in index.values():
        bucket.sort(key=lambda v: -len(v.slots))
    return index


def find_hits(text, index, min_confidence):
    tokens = tokenize(text)
    hits = []
    seen = set()
    for pos, (form, start, _) in enumerate(tokens):
        for var in index.get(form, ()):
            if var.confidence < min_confidence:
                continue
            span = len(var.slots)
            if pos + span > len(tokens):
                continue
            ok = True
            for offset in range(1, span):
                if tokens[pos + offset][0] not in var.slots[offset]:
                    ok = False
                    break
            if not ok:
                continue
            end = tokens[pos + span - 1][2]
            key = (var.rid, start)
            if key in seen:
                continue
            seen.add(key)
            quote = ' '.join(text[max(0, start - 60):end + 60].split())
            hits.append(Hit(var.rid, var.label, var.confidence, start, end, quote))
            break
    return hits


def load_disclaimer(path):
    if not path:
        return None
    lines = Path(path).read_text(encoding='utf-8').splitlines()
    body = [ln.strip() for ln in lines if ln.strip() and not ln.strip().startswith('#')]
    if not body:
        return None
    return norm(' '.join(body))[:80]


def mark_hits(text, hits, disclaimer_core):
    head = norm(text[:HEAD_CHARS])
    for hit in hits:
        left = max(0, hit.start - WINDOW_CHARS)
        window = norm(text[left:hit.end + WINDOW_CHARS])
        found = any(m in window or m in head for m in STATUS_MARKERS)
        if not found and disclaimer_core:
            found = disclaimer_core in window or disclaimer_core in head
        hit.marked = found
    return hits


def read_document(path):
    suffix = path.suffix.lower()
    if suffix in ('.txt', '.md'):
        return path.read_text(encoding='utf-8')
    if suffix == '.docx':
        try:
            import docx
        except ImportError:
            raise SystemExit('для docx нужен python-docx: pip install python-docx')
        document = docx.Document(str(path))
        return '\n'.join(p.text for p in document.paragraphs)
    if suffix == '.pdf':
        try:
            import pdfplumber
        except ImportError:
            raise SystemExit('для pdf нужен pdfplumber: pip install pdfplumber')
        with pdfplumber.open(str(path)) as pdf:
            return '\n'.join(page.extract_text() or '' for page in pdf.pages)
    raise SystemExit(f'формат не поддерживается: {path.name}')


def list_snapshots():
    if not SNAPSHOT_DIR.exists():
        return []
    return sorted(p for p in SNAPSHOT_DIR.glob('*.json'))


def load_snapshot(as_of=None, explicit=None):
    if explicit:
        chosen = Path(explicit)
    else:
        available = list_snapshots()
        if not available:
            raise SystemExit('снимков реестра нет: сначала выполните update')
        chosen = available[-1]
        if as_of:
            fit = [p for p in available if p.stem <= as_of]
            if not fit:
                raise SystemExit(f'нет снимка на дату {as_of} или раньше')
            chosen = fit[-1]
    payload = json.loads(chosen.read_text(encoding='utf-8'))
    records = []
    for i, row in enumerate(payload.get('records', ())):
        records.append(Record(i, row.get('name', ''), row.get('kind', 'person'),
                              row.get('included_on', '')))
    if as_of:
        records = [r for r in records if not r.included_on or r.included_on <= as_of]
    return chosen, payload, records


SKIP_PREFIXES = ('№', 'наимен', 'фамил', 'ф.и.о', 'сведения', 'полное')


def parse_rows(rows):
    records = []
    for row in rows:
        cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
        if not cells:
            continue
        name = cells[0]
        low = norm(name)
        if len(name) < 4 or low.startswith(SKIP_PREFIXES) or low.isdigit():
            continue
        included = ''
        for cell in cells[1:]:
            if len(cell) == 10 and cell[4] in '-.':
                included = cell.replace('.', '-')
        kind = 'org' if len(name.split()) > 3 else 'person'
        records.append({'name': name, 'kind': kind, 'included_on': included})
    return records


def rows_from_file(path):
    suffix = path.suffix.lower()
    if suffix == '.json':
        data = json.loads(path.read_text(encoding='utf-8'))
        items = data.get('records', ()) if isinstance(data, dict) else data
        return [[it.get('name', ''), it.get('included_on', '')] for it in items]
    if suffix == '.csv':
        import csv
        with path.open(encoding='utf-8', newline='') as handle:
            return [row for row in csv.reader(handle, delimiter=';')]
    if suffix == '.xlsx':
        try:
            import openpyxl
        except ImportError:
            raise SystemExit('для xlsx нужен openpyxl: pip install openpyxl')
        book = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        return [list(row) for row in book.active.iter_rows(values_only=True)]
    if suffix == '.docx':
        try:
            import docx
        except ImportError:
            raise SystemExit('для docx нужен python-docx: pip install python-docx')
        document = docx.Document(str(path))
        rows = []
        for table in document.tables:
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
        return rows
    raise SystemExit(f'не умею читать выгрузку: {path.name}')


def cmd_update(args):
    SNAPSHOT_DIR.mkdir(parents=True, exist_ok=True)
    if args.url:
        try:
            import requests
        except ImportError:
            raise SystemExit('для загрузки по ссылке нужен requests: pip install requests')
        response = requests.get(args.url, timeout=60)
        response.raise_for_status()
        suffix = Path(args.url.split('?')[0]).suffix or '.xlsx'
        raw = SNAPSHOT_DIR / ('download' + suffix)
        raw.write_bytes(response.content)
        source = args.url
    elif args.from_file:
        raw = Path(args.from_file)
        source = str(raw)
    else:
        raise SystemExit('укажите --url или --from-file. Адрес и формат выгрузки Минюста менялись, '
                         'вшитый в код URL однажды станет мёртвым и утилита будет молча врать')
    records = parse_rows(rows_from_file(raw))
    previous = list_snapshots()
    if len(records) < MIN_RECORDS and not args.force:
        raise SystemExit(f'разобрано {len(records)} записей при пороге {MIN_RECORDS}: похоже, '
                         'формат выгрузки изменился. Снимок не сохранён')
    if previous and not args.force:
        old = json.loads(previous[-1].read_text(encoding='utf-8'))
        old_count = max(len(old.get('records', ())), 1)
        delta = abs(len(records) - old_count) / old_count
        if delta > MAX_DELTA:
            raise SystemExit(f'расхождение с предыдущим снимком {delta:.0%} при пороге '
                             f'{MAX_DELTA:.0%}. Снимок не сохранён, проверьте выгрузку вручную')
    today = date.today().isoformat()
    payload = {'fetched_on': today, 'source': source, 'count': len(records), 'records': records}
    blob = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    payload['sha256'] = hashlib.sha256(blob.encode('utf-8')).hexdigest()
    target = SNAPSHOT_DIR / f'{today}.json'
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=1), encoding='utf-8')
    MANIFEST.write_text(json.dumps({'latest': target.name, 'count': len(records),
                                    'sha256': payload['sha256']}, ensure_ascii=False, indent=1),
                        encoding='utf-8')
    print(f'снимок {target.name}: записей {len(records)}, sha256 {payload["sha256"][:16]}')
    return 0


def cmd_snapshots(args):
    available = list_snapshots()
    if not available:
        print('снимков нет: выполните update')
        return 0
    for path in available:
        payload = json.loads(path.read_text(encoding='utf-8'))
        digest = payload.get('sha256', '')
        print(f'{path.stem}  записей {payload.get("count", 0)}  sha256 {digest[:16]}')
    return 0


def render_report(path, snapshot, payload, hits, min_confidence):
    count = payload.get('count', len(payload.get('records', ())))
    digest = payload.get('sha256', 'не указан')
    unmarked = [h for h in hits if not h.marked]
    lines = [
        f'# Отчёт вычитки: {path.name}',
        '',
        f'- снимок реестра: {snapshot.name}',
        f'- записей в снимке: {count}',
        f'- sha256 снимка: {digest}',
        f'- порог уверенности: {min_confidence}',
        f'- упоминаний: {len(hits)}, без маркировки рядом: {len(unmarked)}',
        '',
        'Отчёт вспомогательный. Не является юридической консультацией.',
        'Оценку того, относится ли материал к иноагентской деятельности, делает человек.',
        '',
    ]
    if not hits:
        lines.append('Упоминаний не найдено. Это не гарантия: морфология работает правилами, '
                     'редкие и иноязычные фамилии могут быть пропущены. См. README.')
    for hit in sorted(hits, key=lambda h: h.start):
        status = 'маркировка рядом есть' if hit.marked else 'МАРКИРОВКИ РЯДОМ НЕТ'
        lines.extend([
            f'## {hit.label}',
            f'- уверенность: {hit.confidence:.2f}',
            f'- смещение в тексте: {hit.start}',
            f'- {status}',
            f'- контекст: ...{hit.quote}...',
            '',
        ])
    return '\n'.join(lines)


def cmd_scan(args):
    snapshot, payload, records = load_snapshot(args.as_of, args.snapshot)
    index = build_index(records)
    core = load_disclaimer(args.disclaimer_file)
    out_dir = Path(args.out) if args.out else None
    if out_dir:
        out_dir.mkdir(parents=True, exist_ok=True)
    print(f'снимок {snapshot.stem}, записей {len(records)}')
    total_unmarked = 0
    for raw in args.files:
        path = Path(raw)
        text = read_document(path)
        hits = mark_hits(text, find_hits(text, index, args.min_confidence), core)
        unmarked = [h for h in hits if not h.marked]
        total_unmarked += len(unmarked)
        if out_dir:
            report = render_report(path, snapshot, payload, hits, args.min_confidence)
            (out_dir / (path.stem + '.report.md')).write_text(report, encoding='utf-8')
        print(f'{path.name}: упоминаний {len(hits)}, без маркировки {len(unmarked)}')
        for hit in unmarked:
            print(f'  [{hit.confidence:.2f}] {hit.label} @ {hit.start}: ...{hit.quote}...')
    return 1 if total_unmarked else 0


def build_parser():
    parser = argparse.ArgumentParser(description='Оффлайн-вычитка на упоминания иноагентов. '
                                                 'Не является юридической консультацией.')
    sub = parser.add_subparsers(dest='command', required=True)

    upd = sub.add_parser('update', help='сохранить снимок реестра (единственная команда с сетью)')
    upd.add_argument('--url', help='прямая ссылка на выгрузку реестра')
    upd.add_argument('--from-file', help='локальный файл выгрузки: xlsx, csv, json, docx')
    upd.add_argument('--force', action='store_true', help='сохранить снимок несмотря на проверки')
    upd.set_defaults(func=cmd_update)

    snaps = sub.add_parser('snapshots', help='список локальных снимков')
    snaps.set_defaults(func=cmd_snapshots)

    scan = sub.add_parser('scan', help='вычитать документы, оффлайн')
    scan.add_argument('files', nargs='+')
    scan.add_argument('--as-of', dest='as_of', help='состояние реестра на дату, ГГГГ-ММ-ДД')
    scan.add_argument('--snapshot', help='путь к конкретному снимку')
    scan.add_argument('--min-confidence', dest='min_confidence', type=float,
                      default=DEFAULT_MIN_CONFIDENCE)
    scan.add_argument('--out', help='каталог для отчётов')
    scan.add_argument('--disclaimer-file', dest='disclaimer_file',
                      help='файл формы маркировки, например forms/author_2022.txt')
    scan.set_defaults(func=cmd_scan)
    return parser


def main(argv=None):
    args = build_parser().parse_args(argv)
    return args.func(args)


if __name__ == '__main__':
    sys.exit(main())
